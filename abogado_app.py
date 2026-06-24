import streamlit as st
from groq import Groq
from docx import Document
import io
import os
import smtplib
from email.mime.multipart import MIMEMultipart
from email.mime.text import MIMEText
from email.mime.base import MIMEBase
from email import encoders

# Configuración del título principal de la aplicación web
st.title("⚖️ Franc - Asistente Legal Corporativo Experto")
st.write("Sistema inteligente dual para la redacción y análisis de documentos jurídicos blindados en Venezuela.")

# TEXTO CONSTANTE DEL AVISO LEGAL REQUERIDO
AVISO_LEGAL_TEXTO = (
    "⚠️ AVISO LEGAL IMPORTANTE: Este documento ha sido generado mediante Inteligencia Artificial "
    "y constituye únicamente un borrador o modelo de referencia. No sustituye el asesoramiento legal "
    "profesional. Debido a la constante actualización de normativas y criterios discrecionales del SAREN "
    "y tribunales venezolanos, este texto debe ser revisado, adaptado y visado obligatoriamente por un "
    "abogado colegiado (con Inpreabogado) antes de su firma, autenticación o presentación ante cualquier "
    "Notaría o Registro Público/Mercantil. La plataforma y sus desarrolladores no asumen ninguna "
    "responsabilidad por rechazos, daños, perjuicios o consecuencias legales derivadas del uso directo de este documento."
)

# ARCHIVO LOCAL PARA LOG/REGISTRO DEL CONTADOR GLOBAL PERMANENTE
CONTADOR_FILE = "contador_global.txt"

def obtener_y_actualizar_contador():
    """Lee el archivo del contador global, incrementa el valor y lo guarda."""
    if not os.path.exists(CONTADOR_FILE):
        contador = 1
    else:
        with open(CONTADOR_FILE, "r") as f:
            try:
                contador = int(f.read().strip()) + 1
            except ValueError:
                contador = 1
                
    with open(CONTADOR_FILE, "w") as f:
        f.write(str(contador))
        
    return contador

# ------------------------------------------------------------------
# SECCIÓN SECTORIAL: SELECCIÓN DEL EXPERTO IA
# ------------------------------------------------------------------
st.markdown("### 🗂️ Selección de la Especialidad Legal")
experto_seleccionado = st.selectbox(
    "Seleccione el perfil del Abogado Consultor que necesita:",
    [
        "Derecho Inmobiliario y Mercantil (Alquileres, Ventas, Desalojos)",
        "Derecho Mercantil Avanzado (Constitución de Sociedades Anónimas C.A.)"
    ]
)

st.markdown("---")

# ------------------------------------------------------------------
# FASE 0: DINÁMICA SEGÚN EXPERTO SELECCIONADO
# ------------------------------------------------------------------
st.subheader("📋 Fase 0: Datos del Requerimiento Legal")
st.write("Por favor, suministre los siguientes datos para iniciar (copie, pegue y rellene este formato):")

formato_inmueble = """- Tipo de documento/requerimiento: [Ej. Contrato de arrendamiento comercial / Análisis de riesgo de desalojo]
- Datos de las partes: [Ej. Arrendador (Nombre, Cédula/RIF, Domicilio) / Arrendatario (Empresa, Registro Mercantil, Representante)]
- Ubicación y características del inmueble: [Ej. Local comercial de 80mt2 en Mérida, Edo. Mérida. Especificar linderos o estado de pintura/pisos si aplica]
- Uso específico del comercio: [Ej. Carnicería, venta de repuestos, oficinas]
- Moneda y condiciones financieras: [Ej. Canon en USD mensuales, porcentaje para reparaciones menores, indemnización por ocupación indebida]
- Garantías y Pagos: [Ej. Cantidad de meses de depósito, método de pago tasa BCV]
- Duración de la relación y Fecha: [Ej. Contrato por 1 año, Fecha de inicio y fecha de firma]"""

formato_sociedades = """- Oficina de Registro Mercantil de Destino: [Ej. Registro Mercantil de la Circunscripción Judicial del Estado Mérida]
- Denominación Comercial solicitada: [Ej. Inversiones Alfa, C.A. / Corporación Beta, C.A. (Indicar si ya tiene reserva de nombre)]
- Objeto Social: [Ej. Explotación del ramo de restaurantes, compra y venta de víveres, importación de repuestos. Describir actividad principal y conexas]
- Domicilio Social: [Ej. Ciudad de Mérida, Estado Mérida, República Bolivariana de Venezuela]
- Capital Social y Acciones: [Ej. Capital de 500.000,00 bolívares, representado en 500 acciones de valor nominal de 1.000,00 bolívares cada una]
- Identificación de los Accionistas y su Suscripción: [Ej. Accionista A (Nombre, Cédula, 60% de acciones) y Accionista B (Nombre, Cédula, 40% de acciones)]
- Forma de Pago del Capital: [Ej. Mediante inventario de bienes muebles según balance de apertura / Depósito bancario en efectivo]
- Administración y Firma de la Compañía: [Ej. Junta Directiva (Presidente y Vicepresidente), firma conjunta/separada, duración de 5 años]
- Comisario Principal: [Ej. Nombre, Cédula, Licenciado en Contaduría o Administración, con su respectivo Nro. de CPC o CL]"""

if experto_seleccionado == "Derecho Inmobiliario y Mercantil (Alquileres, Ventas, Desalojos)":
    st.code(formato_inmueble, language="text")
else:
    st.code(formato_sociedades, language="text")

datos_usuario = st.text_area("Introduzca aquí los datos recolectados:", height=250)

# ------------------------------------------------------------------
# PROMPTS DE SISTEMA ASIGNADOS POR CASO 
# ------------------------------------------------------------------
prompt_inmuebles = """FASE 0: RECOPILACIÓN REQUERIDA DE INFORMACIÓN (MANDATO INICIAL OBLIGATORIO)
No asumas datos ni comiences a redactar hasta que el usuario te provee las variables requeridas.

ROLE:
Eres un abogado experto en Derecho Inmobiliario y Mercantil en Venezuela, con más de 15 años de experiencia en asesoría corporativa, redacción de contratos complejos y litigios estratégicos. Tu conocimiento abarca el Código Civil venezolano, el Código de Comercio, el Decreto con Rango, Valor y Fuerza de Ley de Regulación del Arrendamiento Inmobiliario para el Uso Comercial (2014), la Sundde, y los criterios vinculantes de la Sala de Casación Civil y Sala Constitucional del Tribunal Supremo de Justicia (TSJ).
 
TASK:
Tu objetivo es brindar asesoría legal de alto nivel, redactar contratos blindados y analizar casos complejos sobre arrendamientos de locales comerciales, compraventas de inmuebles, procedimientos de desalojo y la regulación de bienes inmuebles de uso comercial, corporativo o industrial en Venezuela.
 
PROTOCOLO DE VERACIDAD Y VALIDACIÓN INTERNA (MANDATO ABSOLUTO):
[Se aplican rigurosamente los mandatos de veracidad absoluta y control de calidad]

OUTPUT FORMAT:
- Para Redacción de Contratos: El documento debe ser autocontenido, con un lenguaje formal, reiterativo y conservador, típico de la contratación mercantil venezolana."""

prompt_sociedades = """## FASE 0: RECOPILACIÓN REQUERIDA DE INFORMACIÓN (MANDATO INICIAL OBLIGATORIO)
No asumas datos, nombres, capitales ni comiences a redactar hasta que el usuario te provee las variables requeridas.

## ROLE:
Eres un Abogado Consultor Senior y Especialista en Derecho Mercantil Venezolano, con 30 años de experiencia específica en el diseño, redacción y revisión de Documentos Constitutivos y Estatutos de Compañías Anónimas (C.A.). Tu conocimiento abarca el Código de Comercio venezolano vigente, el Código Civil, la Ley de Registros y del Notariado, las resoluciones y circulares del Servicio Autónomo de Registros y Notarías (SAREN), y la doctrina y criterios vinculantes del Tribunal Supremo de Justicia (TSJ).

## TASK:
Tu objetivo es brindar asesoría legal corporativa de alto nivel, estructurar actas constitutivas blindadas legalmente, redactar estatutos sociales eficientes y analizar la viabilidad e idoneidad de la estructura jurídica de las sociedades comerciales que se someterán a inscripción ante las Oficinas de Registro Mercantil en Venezuela.

## PROTOCOLO DE VERACIDAD Y VALIDACIÓN INTERNA (MANDATO ABSOLUTO):
[Se aplican rigurosamente las reglas del SAREN y el desglose matemático del capital social]

## OUTPUT FORMAT:
- **Para Redacción de Documentos Constitutivos y Estatutos:** El documento debe ser autocontenido, empleando el lenguaje formal, de la tradición mercantil venezolana."""

# ------------------------------------------------------------------
# FUNCIÓN DEL CEREBRO LEGAL (CONEXIÓN CON GROQ CLOUD)
# ------------------------------------------------------------------
def consultar_abogado_ia(datos_fase0, experto):
    try:
        client = Groq(api_key=st.secrets["GROQ_API_KEY"])
        
        if experto == "Derecho Inmobiliario y Mercantil (Alquileres, Ventas, Desalojos)":
            system_prompt = prompt_inmuebles
        else:
            system_prompt = prompt_sociedades
        
        completion = client.chat.completions.create(
            model="llama-3.3-70b-versatile",
            messages=[
                {"role": "system", "content": system_prompt},
                {"role": "user", "content": f"Aquí tienes los datos recolectados de la Fase 0:\n\n{datos_fase0}\n\nPor favor, procede a procesar el requerimiento legal aplicando estrictamente tu protocolo corporativo."}
            ],
            temperature=0.3
        )
        return completion.choices[0].message.content
    except Exception as e:
        return f"Error en el cerebro legal de IA: {e}"

# ------------------------------------------------------------------
# FUNCIÓN PARA CREAR EL ARCHIVO WORD (.DOCX) EN MEMORIA (INCLUYE ADJUNCIÓN DEL AVISO Y CONTADOR)
# ------------------------------------------------------------------
def crear_documento_word(texto_legal, numero_correlativo):
    doc = Document()
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Arial'
    
    for linea in texto_legal.split('\n'):
        doc.add_paragraph(linea)
    
    # ANEXAR CONTROL CORRELATIVO ADMINISTRATIVO
    doc.add_paragraph("\n" + "="*50 + "\n")
    p_control = doc.add_paragraph()
    run_control = p_control.add_run(f"IDENTIFICADOR DE CONTROL PLATAFORMA: N° {numero_correlativo:04d}")
    run_control.bold = True
    
    # ANEXAR EL AVISO LEGAL AL PIE DEL DOCUMENTO GENERADO (.DOCX)
    p_aviso = doc.add_paragraph()
    run_aviso = p_aviso.add_run(AVISO_LEGAL_TEXTO)
    run_aviso.italic = True
    
    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer

# ------------------------------------------------------------------
# FUNCIÓN PARA ENVIAR EL CORREO ELECTRÓNICO CON EL ARCHIVO ADJUNTO INDEXADO
# ------------------------------------------------------------------
def enviar_correo_con_adjunto(correo_destino, texto_legal, archivo_bytes, filename_correlativo):
    try:
        remitente = st.secrets["EMAIL_REMITENTE"]
        password = st.secrets["EMAIL_PASSWORD"]
        
        msg = MIMEMultipart()
        msg['From'] = remitente
        msg['To'] = correo_destino
        msg['Subject'] = f"⚖️ Su Documento Legal Blindado - Control {filename_correlativo.replace('.docx','')}"
        
        cuerpo = f"Hola. Adjunto a este correo encontrará el documento legal generado por el sistema automatizado de IA, registrado bajo el correlativo histórico: {filename_correlativo}."
        msg.attach(MIMEText(cuerpo, 'plain'))
        
        adjunto = MIMEBase('application', 'octet-stream')
        adjunto.set_payload(archivo_bytes.read())
        encoders.encode_base64(adjunto)
        adjunto.add_header('Content-Disposition', 'attachment', filename=filename_correlativo)
        msg.attach(adjunto)
        
        server = smtplib.SMTP('smtp.gmail.com', 587)
        server.starttls()
        server.login(remitente, password)
        server.sendmail(remitente, correo_destino, msg.as_string())
        server.quit()
        return True
    except Exception as e:
        st.error(f"Error al enviar el correo electrónico: {e}")
        return False

# Botón para activar el procesamiento
if st.button("Generar Documento Legal"):
    if datos_usuario.strip() == "":
        st.warning("⚠️ Por favor, debe introducir los datos solicitados en la Fase 0 antes de continuar.")
    else:
        with st.spinner("⚖️ El Abogado de IA está analizando la legislación venezolana y estructurando el requerimiento..."):
            documento_redactado = consultar_abogado_ia(datos_usuario, experto_seleccionado)
            
            # Cálculo del ID Correlativo Global al emitirse con éxito
            id_correlativo = obtener_y_actualizar_contador()
            
            # Guardamos los resultados en la sesión de Streamlit
            st.session_state["documento_resultado"] = documento_redactado
            st.session_state["num_correlativo"] = id_correlativo
            st.success(f"✨ ¡Documento Legal Generado con éxito! Registro Histórico Global: N° {id_correlativo:04d}")

# Si el documento ya existe en memoria, habilitamos las herramientas de entrega
if "documento_resultado" in st.session_state:
    num_id = st.session_state["num_correlativo"]
    nombre_archivo_dinamico = f"documento_legal_{num_id:04d}.docx"
    
    st.markdown(f"### 📄 Previsualización del Documento (Registro de Actividad N° {num_id:04d}):")
    st.write(st.session_state["documento_resultado"])
    
    archivo_word_descarga = crear_documento_word(st.session_state["documento_resultado"], num_id)
    archivo_word_correo = crear_documento_word(st.session_state["documento_resultado"], num_id)
    
    # Opción A: Descarga Local con nombre correlativo numérico
    st.download_button(
        label=f"📥 Descargar Documento: {nombre_archivo_dinamico}",
        data=archivo_word_descarga,
        file_name=nombre_archivo_dinamico,
        mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
    )
    
    st.markdown("---")
    
    # Opción B: Envío por Email Automatizado con Adjunto Indexado
    st.markdown("### 📧 Enviar Documento Directo al Correo")
    email_cliente = st.text_input("Escriba el correo electrónico del destinatario:")
    
    if st.button("Enviar por Correo Electrónico"):
        if email_cliente.strip() == "" or "@" not in email_cliente:
            st.error("⚠️ Por favor, introduzca una dirección de correo electrónico válida.")
        else:
            with st.spinner("🚀 Enviando correo con el archivo adjunto correlativo..."):
                exito = enviar_correo_con_adjunto(email_cliente, st.session_state["documento_resultado"], archivo_word_correo, nombre_archivo_dinamico)
                if exito:
                    st.success(f"📬 ¡Correo enviado con éxito a {email_cliente} conteniendo el archivo {nombre_archivo_dinamico}!")

# ------------------------------------------------------------------
# ANEXAR AVISO LEGAL AL PIE DE LA PÁGINA WEB
# ------------------------------------------------------------------
st.markdown("---")
st.warning(AVISO_LEGAL_TEXTO)

# ------------------------------------------------------------------
# SECCIÓN DE SOPORTE, COMENTARIOS Y CONTACTO
# ------------------------------------------------------------------
st.markdown("---")
st.markdown("### 💡 Sugerencias o Comentarios")
st.write("Para sugerencias o comentarios sobre la aplicación, por favor comuníquese con:")

col1, col2 = st.columns(2)
with col1:
    st.markdown("👤 **Contacto:** Francisco Durán")
    st.markdown("💬 **WhatsApp:** [+58 416-8184675](https://wa.me/584168184675)")
with col2:
    st.markdown("📧 **Email:** [fduranmontillaia@gmail.com](mailto:fduranmontillaia@gmail.com)")
